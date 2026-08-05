# ==============================================================================
#  MODULE 8: EXPORTING TABLES TO EXCEL, WORD, AND LATEX
#  Dataset: CES 2020 (cleaned in Module 1)
#
#  Stack: pandas (Excel/CSV), openpyxl (Excel formatting),
#         tabulate (LaTeX/Markdown), python-docx (Word)
#
#  pip install openpyxl tabulate python-docx
# ==============================================================================

import pandas as pd
import numpy as np
import statsmodels.formula.api as smf
from tabulate import tabulate

ces = pd.read_parquet("CES2020_clean.parquet")


# ==============================================================================
# SECTION 1: BUILD A REGRESSION COEFFICIENT TABLE
# ==============================================================================

# Run three nested models
m1 = smf.ols("imm_restrict ~ education", data=ces).fit()
m2 = smf.ols("imm_restrict ~ education + age + C(sex) + C(census_region)", data=ces).fit()
m3 = smf.ols(
    "imm_restrict ~ education + age + C(sex) + C(census_region) + ideology5 + party_id7",
    data=ces
).fit()

models = {"Bivariate": m1, "Demographics": m2, "Full": m3}


def stars(p):
    if p < 0.001: return "***"
    if p < 0.01:  return "**"
    if p < 0.05:  return "*"
    return ""


def build_coef_table(models_dict):
    """Build a publication-style coefficient table from a dict of fitted models."""
    all_params = sorted(
        set(p for m in models_dict.values() for p in m.params.index),
        key=lambda x: (x == "Intercept", x)
    )

    rows = []
    for param in all_params:
        row = {"term": param}
        for name, m in models_dict.items():
            if param in m.params:
                b  = m.params[param]
                se = m.bse[param]
                p  = m.pvalues[param]
                row[name]          = f"{b:.3f}{stars(p)}"
                row[f"{name}_se"]  = f"({se:.3f})"
            else:
                row[name]         = ""
                row[f"{name}_se"] = ""
        rows.append(row)

    # Add model-level stats
    for stat, fn in [("N", lambda m: int(m.nobs)),
                     ("R²", lambda m: f"{m.rsquared:.3f}"),
                     ("Adj. R²", lambda m: f"{m.rsquared_adj:.3f}")]:
        row = {"term": stat}
        for name, m in models_dict.items():
            row[name]         = fn(m)
            row[f"{name}_se"] = ""
        rows.append(row)

    return pd.DataFrame(rows)


table = build_coef_table(models)
print(table.to_string(index=False))


# ==============================================================================
# SECTION 2: EXPORT TO CSV (UNIVERSAL)
# ==============================================================================

table.to_csv("table_ols_restrict.csv", index=False)
print("Saved: table_ols_restrict.csv")


# ==============================================================================
# SECTION 3: EXPORT TO EXCEL WITH FORMATTING
# ==============================================================================

# openpyxl allows cell-level formatting (bold headers, borders, column widths)

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

wb = openpyxl.Workbook()
ws = wb.active
ws.title = "OLS Regression"

# Write header
ws["A1"] = "OLS Regression: Predictors of Immigration Restrictionism"
ws["A1"].font = Font(bold=True, size=13)
ws.merge_cells("A1:G1")

# Column headers
headers = ["Term", "Bivariate", "(SE)", "Demographics", "(SE)", "Full", "(SE)"]
for col_idx, header in enumerate(headers, start=1):
    cell = ws.cell(row=3, column=col_idx, value=header)
    cell.font      = Font(bold=True)
    cell.alignment = Alignment(horizontal="center")
    cell.fill      = PatternFill(fill_type="solid", fgColor="DDEEFF")

# Write coefficient rows
thin_border = Border(
    bottom=Side(border_style="thin"),
    top=Side(border_style="thin")
)

col_order = ["term",
             "Bivariate", "Bivariate_se",
             "Demographics", "Demographics_se",
             "Full", "Full_se"]

for row_idx, (_, row) in enumerate(table.iterrows(), start=4):
    for col_idx, col_key in enumerate(col_order, start=1):
        cell = ws.cell(row=row_idx, column=col_idx,
                       value=row.get(col_key, ""))
        cell.alignment = Alignment(horizontal="center" if col_idx > 1 else "left")

# Auto-fit column widths (approximate)
for col in ws.columns:
    max_len = max((len(str(cell.value or "")) for cell in col), default=0)
    ws.column_dimensions[col[0].column_letter].width = max(max_len + 2, 10)

# Add a note row
note_row = ws.max_row + 1
ws.cell(row=note_row, column=1,
        value="Note: * p<.05, ** p<.01, *** p<.001. SEs in parentheses.")
ws.cell(row=note_row, column=1).font = Font(italic=True, size=9)

wb.save("table_ols_restrict.xlsx")
print("Saved: table_ols_restrict.xlsx")


# ==============================================================================
# SECTION 4: EXPORT TO WORD WITH PYTHON-DOCX
# ==============================================================================

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

doc = Document()

# Title
title = doc.add_heading("OLS Regression: Predictors of Immigration Restrictionism", level=1)
title.runs[0].font.size = Pt(14)

# Intro paragraph
n = int(m3.nobs)
r2 = m3.rsquared
doc.add_paragraph(
    f"Table 1 presents three nested OLS regression models predicting immigration "
    f"restrictionism. The full model (Model 3) uses N={n:,} observations and "
    f"explains {r2:.1%} of the variance in the outcome."
)

# Build table in docx
col_headers = ["Term", "Bivariate", "Demographics", "Full Model"]
display_cols = ["term", "Bivariate", "Demographics", "Full"]
se_cols      = ["", "Bivariate_se", "Demographics_se", "Full_se"]

word_table = doc.add_table(rows=1, cols=4)
word_table.style = "Table Grid"

# Header row
hdr_cells = word_table.rows[0].cells
for i, h in enumerate(col_headers):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].bold = True

# Data rows (coefficient + SE rows interleaved)
for _, row in table.iterrows():
    coef_row = word_table.add_row().cells
    for i, col in enumerate(display_cols):
        coef_row[i].text = str(row.get(col, ""))

    # SE row (if any SE exists for this row)
    has_se = any(str(row.get(s, "")).strip() not in ("", "nan")
                 for s in se_cols if s)
    if has_se:
        se_row = word_table.add_row().cells
        for i, col in enumerate(se_cols):
            se_row[i].text = str(row.get(col, "")) if col else ""
            se_row[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# Note
note = doc.add_paragraph(
    "Note: * p<.05, ** p<.01, *** p<.001. Robust (HC3) standard errors in parentheses. "
    "Source: CES 2020 (Schaffner, Ansolabehere & Luks, 2021)."
)
note.runs[0].font.size = Pt(9)
note.runs[0].italic    = True

doc.save("table_ols_restrict.docx")
print("Saved: table_ols_restrict.docx")


# ==============================================================================
# SECTION 5: EXPORT TO LATEX
# ==============================================================================

def to_latex_table(models_dict, caption="OLS Regression Results", label="tab:ols"):
    """Build a LaTeX booktabs regression table."""
    col_names = list(models_dict.keys())
    all_params = []
    for name, m in models_dict.items():
        for p in m.params.index:
            if p not in all_params:
                all_params.append(p)

    lines = []
    lines.append(r"\begin{table}[htbp]")
    lines.append(r"\centering")
    lines.append(rf"\caption{{{caption}}}")
    lines.append(rf"\label{{{label}}}")
    lines.append(r"\begin{tabular}{l" + "c" * len(col_names) + "}")
    lines.append(r"\toprule")
    lines.append(" & " + " & ".join(col_names) + r" \\")
    lines.append(r"\midrule")

    for param in all_params:
        row_coef = [param.replace("_", r"\_")]
        row_se   = [""]
        for name, m in models_dict.items():
            if param in m.params:
                b  = m.params[param]
                se = m.bse[param]
                p  = m.pvalues[param]
                row_coef.append(f"{b:.3f}{stars(p)}")
                row_se.append(f"({se:.3f})")
            else:
                row_coef.append("")
                row_se.append("")
        lines.append(" & ".join(row_coef) + r" \\")
        lines.append(" & ".join(row_se)   + r" \\")

    lines.append(r"\midrule")
    # GOF stats
    for stat, fn in [("N", lambda m: str(int(m.nobs))),
                     ("$R^2$", lambda m: f"{m.rsquared:.3f}"),
                     ("Adj.~$R^2$", lambda m: f"{m.rsquared_adj:.3f}")]:
        row = [stat] + [fn(m) for m in models_dict.values()]
        lines.append(" & ".join(row) + r" \\")

    lines.append(r"\bottomrule")
    lines.append(
        r"\multicolumn{" + str(len(col_names) + 1) + r"}{l}{"
        r"\textit{Note:} $^{*}p<.05$, $^{**}p<.01$, $^{***}p<.001$. "
        r"SEs in parentheses.} \\"
    )
    lines.append(r"\end{tabular}")
    lines.append(r"\end{table}")

    return "\n".join(lines)


latex_output = to_latex_table(models, caption="OLS: Immigration Restrictionism")

with open("table_ols_restrict.tex", "w") as f:
    f.write(latex_output)

print("Saved: table_ols_restrict.tex")


# ==============================================================================
# SECTION 6: DESCRIPTIVE STATISTICS TABLE
# ==============================================================================

desc_vars = ["age", "education", "ideology5", "imm_restrict",
             "econ_retro", "self_rated_health", "college", "voted",
             "biden_voter", "white_nh"]

desc = (
    ces[desc_vars]
    .agg(["mean", "std", "median", "min", "max", lambda x: x.notna().sum()])
    .T
    .rename(columns={"<lambda_0>": "N"})
    .round(3)
)

desc.to_csv("table1_descriptives.csv")
print("\nTable 1 (descriptives):")
print(tabulate(desc, headers="keys", tablefmt="pipe", floatfmt=".3f"))


# ==============================================================================
# SECTION 7: LOGIT TABLE WITH ODDS RATIOS
# ==============================================================================

import numpy as np

l1 = smf.logit("voted ~ education + age + C(sex) + ideology5", data=ces).fit(disp=False)
l2 = smf.logit("voted ~ education + age + C(sex) + ideology5 + party_id7 + college", data=ces).fit(disp=False)
l3 = smf.logit("voted ~ education + age + C(sex) + ideology5 + party_id7 + college + white_nh + imm_restrict", data=ces).fit(disp=False)

logit_models = {"Model 1": l1, "Model 2": l2, "Model 3": l3}

def build_or_table(models_dict):
    """Build an odds-ratio table from logit models."""
    all_params = []
    for m in models_dict.values():
        for p in m.params.index:
            if p not in all_params:
                all_params.append(p)

    rows = []
    for param in all_params:
        row = {"term": param}
        for name, m in models_dict.items():
            if param in m.params:
                OR = np.exp(m.params[param])
                se = m.bse[param]
                p  = m.pvalues[param]
                ci_lo = np.exp(m.conf_int().loc[param, 0])
                ci_hi = np.exp(m.conf_int().loc[param, 1])
                row[name] = f"{OR:.3f}{stars(p)}"
                row[f"{name}_ci"] = f"[{ci_lo:.3f}, {ci_hi:.3f}]"
            else:
                row[name]         = ""
                row[f"{name}_ci"] = ""
        rows.append(row)

    for stat, fn in [("N", lambda m: int(m.nobs)),
                     ("AIC", lambda m: f"{m.aic:.1f}"),
                     ("Pseudo R²", lambda m: f"{1 - m.llf / smf.logit('voted~1', data=ces).fit(disp=False).llf:.3f}")]:
        row = {"term": stat}
        for name, m in models_dict.items():
            row[name]         = fn(m)
            row[f"{name}_ci"] = ""
        rows.append(row)

    return pd.DataFrame(rows)

or_table = build_or_table(logit_models)
or_table.to_csv("table_logit_turnout.csv", index=False)
print("\nLogit OR Table saved: table_logit_turnout.csv")


# ==============================================================================
# SECTION 8: EXPORT MULTIPLE SHEETS TO EXCEL
# ==============================================================================

# pd.ExcelWriter allows writing multiple DataFrames to separate sheets
with pd.ExcelWriter("ces_analysis_output.xlsx", engine="openpyxl") as writer:
    desc.to_excel(writer, sheet_name="Descriptive Statistics")
    table.to_excel(writer, sheet_name="OLS Regression", index=False)
    or_table.to_excel(writer, sheet_name="Logit ORs", index=False)

print("Saved: ces_analysis_output.xlsx (3 sheets)")


# ==============================================================================
# SECTION 9: MARKDOWN TABLE (FOR GITHUB README OR NOTEBOOKS)
# ==============================================================================

print("\nMarkdown table (for README or notebook):")
print(tabulate(
    desc.head(6),
    headers="keys",
    tablefmt="github",
    floatfmt=".2f"
))

print("\nModule 8 complete.")
